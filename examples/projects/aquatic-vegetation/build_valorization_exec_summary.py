from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Austin_Lake_Vegetation_Valorization_Executive_Summary_2026-06-17.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "111111"
MUTED = "555555"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=8.6, color=INK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in [("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")]:
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), val)
                node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4.2)
    normal.paragraph_format.line_spacing = 1.04

    for name, size, color, before, after in [
        ("Heading 1", 14.5, BLUE, 8, 4),
        ("Heading 2", 11.6, BLUE, 6, 3),
        ("Heading 3", 10.2, DARK_BLUE, 4, 2),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Austin Lake Vegetation Valorization | Executive Summary | 2026-06-17")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def para(doc, text="", style=None, bold_prefix=None, size=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Calibri"
        if size:
            r.font.size = Pt(size)
        rest = text[len(bold_prefix):]
        if rest:
            r2 = p.add_run(rest)
            r2.font.name = "Calibri"
            if size:
                r2.font.size = Pt(size)
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
        if size:
            r.font.size = Pt(size)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.24 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9.4)
    return p


def callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(9.8)
    r1.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r2 = p.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(9.4)
    return table


def simple_table(doc, headers, rows, widths, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=font_size, color=DARK_BLUE)
        set_cell_shading(cell, LIGHT_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=font_size)
    return table


def add_source_paragraph(doc, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    add_hyperlink(p, url, url)
    for run in p.runs:
        run.font.size = Pt(8)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document()
    style_doc(doc)

    # Page 1
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Austin Lake Vegetation Valorization")
    r.font.name = "Calibri"
    r.font.size = Pt(23)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    sr = subtitle.add_run("Four-page executive summary for partner discussion | Prepared June 17, 2026")
    sr.font.name = "Calibri"
    sr.font.size = Pt(9.5)
    sr.font.color.rgb = RGBColor.from_string(MUTED)

    callout(
        doc,
        "Recommendation:",
        "Proceed with a clean-macrophyte resource-recovery pilot, not a consumer-product extraction venture. The near-term revenue stream is a mix of customer-paid recovery fees, avoided disposal, and partner processing into compost or anaerobic digestion feedstock. Food, feed, supplements, and wild-algae extracts should remain off the table until controlled sourcing and safety testing are proven.",
    )
    para(doc, "What the second revenue stream is", style="Heading 1")
    para(
        doc,
        "The collected biomass can create value, but the valuable unit is not the plant itself. It is the compliance-managed chain of custody around a wet, regulated, seasonal waste stream. The startup should charge for removal first, then capture incremental margin by routing clean Hydrilla and related submerged vegetation away from landfill into a tested compost, digestion, or soil-amendment pathway.",
    )
    para(
        doc,
        "The Austin context matters. Lake Austin has a documented 2026 Hydrilla spike, but it is also a drinking-water reservoir where herbicides are not allowed and where mechanical work must avoid spreading fragments. Lady Bird Lake and parts of Lake Austin also have recurring potentially toxic blue-green algae. That means clean macrophytes and suspected cyanobacteria must be separated from the first day of operations.",
    )
    simple_table(
        doc,
        ["Pathway", "Partner-ready?", "Commercial read", "First action"],
        [
            ["Compost / soil amendment", "High", "Best first path; monetizes avoided disposal and environmental reporting more than product sales.", "Pilot with a permitted composter; test metals, nutrients, moisture, and toxins."],
            ["Anaerobic digestion", "Medium", "Technically plausible but energy value per wet ton is low; partner gate economics decide.", "Run BMP test and ask Hornsby Bend/private digesters about co-digestion intake."],
            ["Biochar / hydrochar", "Medium-low", "Potential premium soil product, but drying wet biomass can erase margin.", "Only test after compost/AD partner route works."],
            ["Feed, supplement, extracts", "Low", "Not feasible from wild Austin lake biomass at partner-ready risk level.", "Do not pursue in first 24 months."],
        ],
        [1950, 1450, 4300, 1660],
        font_size=7.3,
    )
    p = para(
        doc,
        "Bottom line: the business should sell a clean removal service plus documented resource recovery. Product revenue is upside; avoided disposal and recovery fees are the bankable economics.",
    )
    p.paragraph_format.space_before = Pt(8)
    page_break(doc)

    # Page 2
    para(doc, "Operating Model And Cost Stack", style="Heading 1")
    para(
        doc,
        "Start with a partner-processing model. Buying or building a composting, digestion, or char facility too early would turn a service startup into a regulated waste-processing company before feedstock quality and volume are known. The first year should prove that vegetation can be segregated, permitted, hauled, tested, accepted by a processor, and tied to a customer-paid recovery fee.",
    )
    para(doc, "Required workflow", style="Heading 2")
    bullet(doc, "Collect only permitted clean macrophytes for reuse. Treat suspected blue-green algae as a separate hazard stream.")
    bullet(doc, "Dewater at the jobsite or staging point where allowed; hauling water is the largest avoidable cost.")
    bullet(doc, "Use covered, contained transport to prevent Hydrilla or milfoil fragment spread.")
    bullet(doc, "Run a minimum lab panel: moisture, solids, NPK, C:N, metals, and cyanotoxins when any algae or mixed mats are possible.")
    bullet(doc, "Route clean loads to a composter or digestion partner; route contaminated or uncertain loads to disposal.")
    para(doc, "Setup and variable-cost assumptions", style="Heading 2")
    simple_table(
        doc,
        ["Cost item", "Pilot range", "What it covers"],
        [
            ["Permits, SOPs, insurance review", "$8k-$18k", "TPWD treatment proposal/permit workflow, transport protocol, customer authorizations, safety rules."],
            ["Testing program", "$4k-$12k", "TAMU nutrient/solids panels are inexpensive; cyanotoxin and metals panels drive cost when algae risk is present."],
            ["Transport setup", "$10k-$35k", "Tarps, containment, pumps, geotextile bags, bins, scale process, or leased watertight trailers."],
            ["Hauling and gate fees", "$25-$70/wet ton", "Depends on distance, dewatering, load size, and whether material is accepted as compostable organics or landfill waste."],
            ["Processing partner", "$0-$35/wet ton", "A composter or digester may charge, discount, or reject depending on contamination, moisture, and process fit."],
        ],
        [2300, 1700, 5360],
        font_size=7.8,
    )
    para(
        doc,
        "Local anchors support these assumptions. TDS posts municipal solid waste landfill rates at $60/ton plus fees, and brush/leaves/untreated lumber at $30/ton for large loads at the Creedmoor landfill. Austin Water's Dillo Dirt program proves local compost infrastructure exists, but its feedstock is controlled biosolids plus brush, not automatically invasive aquatic weeds. Partner acceptance must be negotiated.",
    )
    para(doc, "Why product extraction is not the near-term plan", style="Heading 2")
    para(
        doc,
        "Hydrilla is roughly 95% water when harvested. That makes drying, milling, extracting, and packaging expensive. The energy value in anaerobic digestion is also modest per wet ton: using a conservative 5% dry solids assumption, one wet ton may only hold enough volatile solids for roughly 7-11 cubic meters of methane before process losses. At commodity energy prices, this is not enough to justify collection; it only improves a removal business that is already being paid.",
    )
    page_break(doc)

    # Page 3
    para(doc, "Illustrative ROI Model", style="Heading 1")
    para(
        doc,
        "The model below isolates the second revenue stream. It assumes the removal business is already paid for cutting and clearing vegetation. Valorization adds margin only when customers pay for resource recovery, landfill or disposal cost is avoided, or a processor pays/discounts for clean feedstock. These are planning cases, not quotes.",
    )
    simple_table(
        doc,
        ["Scenario", "Wet tons/year", "Value captured/wet ton", "Variable cost/wet ton", "Fixed setup", "Year-1 incremental profit", "Year-2 run-rate"],
        [
            ["Pilot", "300", "$35", "$25", "$18k", "-$15k", "$3k"],
            ["Base", "1,000", "$60", "$28", "$30k", "$2k", "$32k"],
            ["Strong", "2,500", "$85", "$32", "$65k", "$67.5k", "$132.5k"],
        ],
        [1200, 1150, 1450, 1450, 1300, 1450, 1360],
        font_size=7.4,
    )
    para(doc, "Calculation example", style="Heading 2")
    para(
        doc,
        "Base case: 1,000 wet tons x ($60 value captured - $28 variable handling cost) = $32,000 contribution. Subtract $30,000 of first-year setup and testing, and year-one incremental profit is approximately $2,000. In year two, if the same route is repeatable and setup costs do not recur, the run-rate contribution is about $32,000. The strong case reaches more than $67,000 in first-year incremental profit because tonnage spreads fixed costs and customers pay for documented resource recovery.",
    )
    simple_table(
        doc,
        ["Value component", "Pilot", "Base", "Strong"],
        [
            ["Avoided disposal vs landfill/uncertain waste route", "$10/t", "$25/t", "$35/t"],
            ["Customer resource-recovery fee", "$20/t", "$30/t", "$40/t"],
            ["Processor/product credit", "$5/t", "$5/t", "$10/t"],
            ["Gross value captured", "$35/t", "$60/t", "$85/t"],
        ],
        [3600, 1920, 1920, 1920],
        font_size=7.9,
    )
    para(doc, "Combined business effect", style="Heading 2")
    para(
        doc,
        "The main removal business still drives the economics. If 1,000 wet tons corresponds to roughly 100-140 customer jobs and $300k-$450k in removal revenue, a 35%-45% removal gross margin produces $105k-$200k of gross profit before valorization. The base valorization case adds about $32k of annual run-rate contribution, or roughly 7-11 percentage points of gross profit on a $300k-$450k removal business. This is meaningful, but it is not enough to carry the company without paid removal demand.",
    )
    para(
        doc,
        "Sensitivity: hauling distance and water content dominate margin. A $15/wet ton increase in transport/processing cost wipes out almost half of the base-case contribution. Conversely, a $20/wet ton recovery fee paid by HOAs, marinas, or municipalities can turn a break-even disposal pathway into a profit center.",
    )
    page_break(doc)

    # Page 4
    para(doc, "Partner-Ready Plan", style="Heading 1")
    para(doc, "Decision", style="Heading 2")
    para(
        doc,
        "Build the second revenue stream as a disciplined pilot attached to the existing removal wedge. The right partner ask is not 'buy our lake weeds.' It is 'help us convert a recurring regulated waste stream into a tested local soil, compost, or digestion feedstock while documenting nutrient removal and avoiding landfill.'",
    )
    para(doc, "90-day pilot", style="Heading 2")
    bullet(doc, "Secure written regulatory guidance for collection, transport, and reuse of Hydrilla-dominant biomass.")
    bullet(doc, "Sign one processor MOU with a composter or digester and one backup disposal route.")
    bullet(doc, "Run 5-10 test loads with chain-of-custody photos, tonnage, moisture, lab panels, haul cost, gate outcome, and customer pricing.")
    bullet(doc, "Offer customers a resource-recovery line item: $25-$75 per wet ton or a fixed $250-$1,000 site fee for documented compliant diversion.")
    bullet(doc, "Do not co-mingle suspected cyanobacteria; pause, test, and dispose/treat separately.")
    para(doc, "Go / hold / no-go", style="Heading 2")
    simple_table(
        doc,
        ["Decision", "Threshold"],
        [
            ["Go", "At least 500 clean wet tons/year under customer-paid recovery fees; accepted by a processor; net contribution above $25/wet ton."],
            ["Hold", "Processor accepts material but only at landfill-equivalent cost; customer recovery fee unproven; lab data still variable."],
            ["No-go", "Cyanotoxin/metals risk is frequent, loads cannot be segregated, or transport/processing cost exceeds value by more than $20/wet ton."],
        ],
        [1500, 7860],
        font_size=7.8,
    )
    para(doc, "Risks to disclose", style="Heading 2")
    para(
        doc,
        "The venture should be candid that this is not yet a validated product business. The largest risks are regulatory permission, processor acceptance, wet hauling cost, harmful-algae contamination, and the possibility that product revenue is negligible. The upside is strongest when the service is already paid, customers value compliant diversion, and processor partners can use clean biomass without major process changes.",
    )
    para(doc, "Key source notes", style="Heading 2")
    add_source_paragraph(doc, "City of Austin Lake Austin and algae context", "https://www.austintexas.gov/watershed-protection/projects/lake-austin")
    add_source_paragraph(doc, "TPWD nuisance aquatic vegetation and permit rules", "https://tpwd.texas.gov/landwater/water/environconcerns/nuisance_plants/")
    add_source_paragraph(doc, "TDS gate rates and composting infrastructure", "https://www.texasdisposal.com/gate-rates/")
    add_source_paragraph(doc, "Austin Water Dillo Dirt / Hornsby Bend composting and digestion", "https://www.austintexas.gov/water/programs/dillo-dirt")
    add_source_paragraph(doc, "Texas A&M AgriLife testing lab pricing forms", "https://soiltesting.tamu.edu/sample-page/submittal-forms/")
    add_source_paragraph(doc, "USACE mechanical harvesting fresh-weight assumptions", "https://corpslakes.erdc.dren.mil/employees/invasive/pdfs/MechanicalHarvesting.pdf")
    add_source_paragraph(doc, "EPA AD tipping-fee and co-digestion context", "https://www.epa.gov/anaerobic-digestion/anaerobic-digestion-facilities-processing-food-waste-us-2020-2021")
    add_source_paragraph(doc, "Texas composting general requirements", "https://www.law.cornell.edu/regulations/texas/30-Tex-Admin-Code-SS-332-4")
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
